import os
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_gost_document(output_path):
    doc = docx.Document()
    
    # Настройка стилей по ГОСТу
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Настройка полей
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Список файлов в нужном порядке
    files_to_compile = [
        "Структура_и_Оглавление.md",
        "0_Введение.md",
        "1_Глава_Теория.md",
        "2_Глава_Методология.md",
        "3_Глава_Практика.md",
        "4_Заключение.md",
        "5_Приложения.md"
    ]
    
    base_dir = os.path.join("..", "Текст_работы")
    
    for filename in files_to_compile:
        filepath = os.path.join(base_dir, filename)
        if not os.path.exists(filepath):
            print(f"Пропущен: {filepath}")
            continue
            
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Убираем часть markdown разметки и добавляем в word
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            p = doc.add_paragraph()
            
            # Настройка интервала и абзаца (ГОСТ)
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Обработка заголовков
            if line.startswith('#'):
                header_text = re.sub(r'^#+\s*', '', line)
                run = p.add_run(header_text)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
            else:
                # Обработка жирного текста (простая)
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
                        
        p = doc.add_paragraph() # Разрыв между файлами
        # doc.add_page_break() # Можно добавить разрыв страницы для каждой главы

    doc.save(output_path)
    print(f"Диплом (сборка ГОСТ) успешно сохранен: {output_path}")

if __name__ == "__main__":
    out_dir = os.path.join("..", "Готовый_Диплом")
    os.makedirs(out_dir, exist_ok=True)
    create_gost_document(os.path.join(out_dir, "Дипломная_работа_ГОСТ_ФИНАЛ.docx"))
